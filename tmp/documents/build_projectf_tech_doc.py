from __future__ import annotations

import math
import sys
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Git\ProjectF")
OUT_DIR = ROOT / "output" / "documents"
TMP_DIR = ROOT / "tmp" / "documents"
OUT_PATH = OUT_DIR / "ProjectF_기술소개서.docx"
ARCH_IMG = TMP_DIR / "projectf_architecture.png"
CONVEYOR_IMG = TMP_DIR / "projectf_conveyor_pipeline.png"
RAIL_IMG = TMP_DIR / "projectf_rail_pipeline.png"

TABLE_HELPER_DIR = Path(
    r"C:\Users\dk601\.cache\codex-runtimes\codex-primary-runtime\plugins"
    r"\openai-primary-runtime\plugins\documents\skills\documents\scripts"
)
sys.path.insert(0, str(TABLE_HELPER_DIR))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


FONT_NAME = "Malgun Gothic"
MONO_FONT_NAME = "Consolas"
NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "17212B"
MUTED = "5B6573"
LIGHT = "E8EEF5"
PALE = "F4F6F9"
WHITE = "FFFFFF"
LINE = "CCD6E0"
SUCCESS = "1E6B52"
WARNING = "8A5A00"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_border(cell, color: str = LINE, size: int = 6) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def set_paragraph_keep(paragraph, *, keep_with_next=False, keep_together=False) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if keep_with_next:
        p_pr.append(OxmlElement("w:keepNext"))
    if keep_together:
        p_pr.append(OxmlElement("w:keepLines"))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("PROJECTF  ·  ")
    run.font.name = FONT_NAME
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, value, end):
        run._r.append(node)


def add_horizontal_rule(paragraph, color: str = BLUE, size: int = 12) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_custom_numbering(document: Document, kind: str) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl.append(lvl_text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    lvl.append(suffix)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(numbering.index(first_num), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.1


def set_run_font(run, *, size=None, color=None, bold=None, name=FONT_NAME, italic=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size, color, before, after in (
        ("Title", 28, NAVY, 0, 10),
        ("Subtitle", 12, MUTED, 0, 12),
        ("Heading 1", 15, BLUE, 14, 7),
        ("Heading 2", 11.5, BLUE, 9, 4),
        ("Heading 3", 10.5, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Lead" not in styles:
        lead = styles.add_style("Lead", WD_STYLE_TYPE.PARAGRAPH)
    else:
        lead = styles["Lead"]
    lead.font.name = FONT_NAME
    lead._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    lead.font.size = Pt(10.8)
    lead.font.color.rgb = rgb(DARK_BLUE)
    lead.paragraph_format.space_after = Pt(7)
    lead.paragraph_format.line_spacing = 1.2

    if "Caption PF" not in styles:
        caption = styles.add_style("Caption PF", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles["Caption PF"]
    caption.font.name = FONT_NAME
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    caption.font.size = Pt(7.8)
    caption.font.color.rgb = rgb(MUTED)
    caption.font.italic = True
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(5)

    if "Code PF" not in styles:
        code = styles.add_style("Code PF", WD_STYLE_TYPE.CHARACTER)
    else:
        code = styles["Code PF"]
    code.font.name = MONO_FONT_NAME
    code._element.rPr.rFonts.set(qn("w:eastAsia"), MONO_FONT_NAME)
    code.font.size = Pt(8.5)
    code.font.color.rgb = rgb(DARK_BLUE)


def add_heading(document: Document, text: str, level: int = 1, kicker: str | None = None):
    if kicker:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(kicker.upper())
        set_run_font(run, size=8, color=MUTED, bold=True)
        set_paragraph_keep(p, keep_with_next=True)
    heading = document.add_heading(text, level=level)
    return heading


def add_bullet(document: Document, text: str, num_id: int, bold_prefix: str | None = None):
    p = document.add_paragraph()
    apply_numbering(p, num_id)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True, color=DARK_BLUE)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_numbered(document: Document, text: str, num_id: int):
    p = document.add_paragraph()
    apply_numbering(p, num_id)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_callout(document: Document, title: str, body: str, *, accent: str = BLUE):
    table = document.add_table(rows=1, cols=1)
    table.allow_autofit = False
    cell = table.cell(0, 0)
    prevent_row_split(table.rows[0])
    set_cell_shading(cell, PALE)
    set_cell_border(cell, color=accent, size=10)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=9.5, color=accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.08
    r2 = p2.add_run(body)
    set_run_font(r2, size=8.6, color=INK)
    apply_table_geometry(table, [9360])
    return table


def add_metadata_table(document: Document, rows: list[tuple[str, str]]):
    table = document.add_table(rows=len(rows), cols=2)
    for idx, (label, value) in enumerate(rows):
        prevent_row_split(table.rows[idx])
        left, right = table.rows[idx].cells
        set_cell_shading(left, LIGHT)
        for cell in (left, right):
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        lp = left.paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        lr = lp.add_run(label)
        set_run_font(lr, size=8.2, color=DARK_BLUE, bold=True)
        rp = right.paragraphs[0]
        rp.paragraph_format.space_after = Pt(0)
        rr = rp.add_run(value)
        set_run_font(rr, size=8.7, color=INK)
    apply_table_geometry(table, column_widths_from_weights([1.15, 4.85], 9360))
    return table


def add_matrix_table(document: Document, headers: list[str], rows: list[list[str]], weights: list[float]):
    table = document.add_table(rows=1, cols=len(headers))
    header = table.rows[0]
    set_repeat_table_header(header)
    prevent_row_split(header)
    for i, text in enumerate(headers):
        cell = header.cells[i]
        set_cell_shading(cell, LIGHT)
        set_cell_border(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=7.8, color=DARK_BLUE, bold=True)
    for row_values in rows:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for i, value in enumerate(row_values):
            cell = cells[i]
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, size=7.8, color=INK)
    apply_table_geometry(table, column_widths_from_weights(weights, 9360))
    return table


def add_code_path(document: Document, label: str, path: str):
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{label}  ")
    set_run_font(r1, size=8.5, color=MUTED, bold=True)
    r2 = p.add_run(path)
    r2.style = "Code PF"
    return p


def load_fonts():
    font_path = Path(r"C:\Windows\Fonts\malgun.ttf")
    bold_path = Path(r"C:\Windows\Fonts\malgunbd.ttf")
    if not bold_path.exists():
        bold_path = font_path
    return {
        "title": ImageFont.truetype(str(bold_path), 48),
        "h": ImageFont.truetype(str(bold_path), 30),
        "body": ImageFont.truetype(str(font_path), 23),
        "small": ImageFont.truetype(str(font_path), 19),
    }


def rounded_box(draw, xy, fill, outline, title, body, fonts, *, title_color=NAVY, body_color=INK):
    draw.rounded_rectangle(xy, radius=22, fill=f"#{fill}", outline=f"#{outline}", width=3)
    x1, y1, x2, y2 = xy
    draw.text((x1 + 28, y1 + 20), title, font=fonts["h"], fill=f"#{title_color}")
    lines = body.split("\n")
    y = y1 + 72
    for line in lines:
        draw.text((x1 + 28, y), line, font=fonts["small"], fill=f"#{body_color}")
        y += 31


def arrow(draw, start, end, color=BLUE, width=6):
    draw.line([start, end], fill=f"#{color}", width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 18
    for delta in (2.55, -2.55):
        point = (
            end[0] + length * math.cos(angle + delta),
            end[1] + length * math.sin(angle + delta),
        )
        draw.line([end, point], fill=f"#{color}", width=width)


def make_architecture_diagram(path: Path) -> None:
    fonts = load_fonts()
    img = Image.new("RGB", (1800, 930), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 40), "ProjectF 상태·시뮬레이션·표현 아키텍처", font=fonts["title"], fill=f"#{NAVY}")
    rounded_box(d, (80, 150, 520, 350), LIGHT, LINE, "입력·게임플레이", "설치 / 제작 / 운송\n플레이어·UI 명령", fonts)
    rounded_box(d, (680, 150, 1120, 350), PALE, BLUE, "라이브 런타임", "로드된 청크의 설치물\n스케줄 기반 Tick 처리", fonts)
    rounded_box(d, (1280, 150, 1720, 350), LIGHT, LINE, "표현 계층", "가상 렌더 배치\nBurst Transform Job / BRG", fonts)
    rounded_box(d, (380, 560, 820, 790), PALE, DARK_BLUE, "권위 상태", "BlockStateStore\n컨베이어·설치물·청크 상태", fonts)
    rounded_box(d, (980, 560, 1420, 790), LIGHT, LINE, "비활성 영역", "Background Simulator\n저장 상태 기반 시간 진행", fonts)
    arrow(d, (520, 250), (680, 250))
    arrow(d, (1120, 250), (1280, 250))
    arrow(d, (890, 350), (690, 560), DARK_BLUE)
    arrow(d, (820, 675), (980, 675), DARK_BLUE)
    arrow(d, (980, 720), (820, 720), DARK_BLUE)
    d.text((534, 205), "명령", font=fonts["small"], fill=f"#{MUTED}")
    d.text((1145, 205), "스냅샷", font=fonts["small"], fill=f"#{MUTED}")
    d.text((705, 430), "언로드 시 캡처 / 로드 시 복원", font=fonts["small"], fill=f"#{MUTED}")
    d.text((825, 625), "시간 진행", font=fonts["small"], fill=f"#{MUTED}")
    img.save(path, quality=95)


def make_conveyor_diagram(path: Path) -> None:
    fonts = load_fonts()
    img = Image.new("RGB", (1800, 650), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 38), "컨베이어 처리 파이프라인", font=fonts["title"], fill=f"#{NAVY}")
    boxes = [
        ("Tick Scheduler", "간격별 버킷\n쿼터·커서 분산"),
        ("Lane State", "아이템 위치\n연결·진행 상태"),
        ("Transform Job", "NativeArray 재사용\nIJobParallelFor"),
        ("Spatial Batch", "Mesh·Material·Cell\nBatch Key 구성"),
        ("BRG / Fallback", "GPU 버퍼 업로드\n호환성 기반 전환"),
    ]
    x = 45
    for i, (title, body) in enumerate(boxes):
        w = 300
        rounded_box(d, (x, 185, x + w, 445), PALE if i % 2 else LIGHT, BLUE if i in (2, 4) else LINE, title, body, fonts)
        if i < len(boxes) - 1:
            arrow(d, (x + w, 315), (x + w + 55, 315))
        x += 355
    d.text((65, 520), "핵심: 논리 이동과 시각 표현을 분리하고, 반복 할당 대신 재사용 가능한 버퍼와 배치를 사용한다.", font=fonts["body"], fill=f"#{DARK_BLUE}")
    img.save(path, quality=95)


def make_rail_diagram(path: Path) -> None:
    fonts = load_fonts()
    img = Image.new("RGB", (1800, 670), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 38), "철도 자동 운행: 경로 데이터에서 상태 머신까지", font=fonts["title"], fill=f"#{NAVY}")
    boxes = [
        ("Rail Sampling", "중심선·접선\n누적 거리 캐시"),
        ("Graph Build", "노드·간선 구성\n버전 기반 무효화"),
        ("Route Search", "우선순위 큐\n역주행·급회전 비용"),
        ("AutoDrive FSM", "Planning / Moving\nDocking / Waiting"),
    ]
    x = 80
    for i, (title, body) in enumerate(boxes):
        rounded_box(d, (x, 185, x + 350, 450), PALE if i % 2 else LIGHT, BLUE if i in (1, 2) else LINE, title, body, fonts)
        if i < len(boxes) - 1:
            arrow(d, (x + 350, 317), (x + 425, 317))
        x += 425
    d.text((80, 525), "노선 변경 → 그래프 버전 증가 → 캐시 무효화 → 경로 재계산 → 운행 상태 전이", font=fonts["body"], fill=f"#{DARK_BLUE}")
    img.save(path, quality=95)


def add_picture(document: Document, path: Path, caption: str, width=6.35):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = document.add_paragraph(caption, style="Caption PF")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def page_break(document: Document):
    document.add_page_break()


def build_document() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    TMP_DIR.mkdir(parents=True, exist_ok=True)
    make_architecture_diagram(ARCH_IMG)
    make_conveyor_diagram(CONVEYOR_IMG)
    make_rail_diagram(RAIL_IMG)

    doc = Document()
    doc.core_properties.title = "ProjectF 기술 소개서"
    doc.core_properties.subject = "Unity 기반 자동화 시뮬레이션 프로젝트 기술 포트폴리오"
    doc.core_properties.author = "한택근"
    doc.core_properties.keywords = "Unity, C#, Conveyor, BatchRendererGroup, Burst, Railway, Chunk Streaming, Save"
    doc.core_properties.comments = "Repository snapshot: 2026-09-01"

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    configure_styles(doc)
    bullet_num = add_custom_numbering(doc, "bullet")
    overview_num = add_custom_numbering(doc, "decimal")
    streaming_num = add_custom_numbering(doc, "decimal")
    ai_num = add_custom_numbering(doc, "decimal")
    next_num = add_custom_numbering(doc, "decimal")

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("PROJECTF  /  TECHNICAL PORTFOLIO")
    set_run_font(hr, size=8, color=MUTED, bold=True)
    add_horizontal_rule(hp, color=LINE, size=5)
    add_page_number(section.footer.paragraphs[0])

    # Cover
    for _ in range(4):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kr = kicker.add_run("UNITY FACTORY AUTOMATION SIMULATION")
    set_run_font(kr, size=9, color=BLUE, bold=True)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("ProjectF")
    set_run_font(tr, size=34, color=NAVY, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    sr = subtitle.add_run("기술 소개서")
    set_run_font(sr, size=20, color=BLUE, bold=True)
    line = doc.add_paragraph()
    line.paragraph_format.left_indent = Inches(1.55)
    line.paragraph_format.right_indent = Inches(1.55)
    add_horizontal_rule(line, color=BLUE, size=10)
    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline.paragraph_format.space_before = Pt(14)
    tagline.paragraph_format.space_after = Pt(18)
    r = tagline.add_run("대규모 자동화 시뮬레이션을 위한\n상태 분리 · 배치 처리 · 스트리밍 설계")
    set_run_font(r, size=15, color=DARK_BLUE, bold=True)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.line_spacing = 1.5
    r = meta.add_run("한택근  |  Client Programmer\n개인 프로젝트  |  Repository Snapshot 2026.09\nUnity 6000.4 · C# · URP · Burst · BatchRendererGroup")
    set_run_font(r, size=9.5, color=MUTED)
    doc.add_paragraph()
    add_callout(
        doc,
        "핵심 기술 명제",
        "모든 오브젝트를 항상 GameObject로 유지하지 않는다. 논리 상태를 권위 데이터로 두고, 로드 범위·백그라운드 시뮬레이션·가상 렌더링이 같은 상태를 각자의 비용 구조에 맞게 소비하도록 설계했다.",
        accent=BLUE,
    )

    # Page 2: Overview
    page_break(doc)
    add_heading(doc, "프로젝트 정의", 1, "01 / OVERVIEW")
    lead = doc.add_paragraph(style="Lead")
    lead.add_run(
        "ProjectF는 자원 채취, 가공, 운송, 자동화, 영역 확장을 하나의 연속된 시뮬레이션으로 연결하는 Unity 기반 팩토리 자동화 프로젝트다."
    )
    add_metadata_table(
        doc,
        [
            ("개발 형태", "1인 개발 · 게임플레이/시스템/도구 전반 직접 구현"),
            ("엔진", "Unity 6000.4.0f1 · Universal Render Pipeline 17.4.0"),
            ("핵심 패키지", "Burst 1.8.28 · Collections 6.4.0 · Mathematics 1.3.3"),
            ("저장 포맷", "PF_SAVE 매직 헤더를 사용하는 버전형 바이너리 포맷 · 현재 버전 43"),
        ],
    )
    add_heading(doc, "해결해야 했던 세 가지 구조적 문제", 2)
    add_numbered(doc, "컨베이어·설치물 수가 늘어날수록 Tick, Transform 계산, Draw Call 비용이 함께 증가한다.", overview_num)
    add_numbered(doc, "플레이어가 보지 않는 청크에서도 생산과 운송의 논리적 시간은 계속 진행되어야 한다.", overview_num)
    add_numbered(doc, "실시간 상태, 언로드된 상태, 저장 파일의 상태가 중복되거나 소실되지 않아야 한다.", overview_num)
    add_heading(doc, "설계 방향", 2)
    add_matrix_table(
        doc,
        ["문제 축", "선택한 접근", "대표 구현"],
        [
            ["업데이트 비용", "개별 Update 대신 중앙 스케줄링", "MapObjectTickManager"],
            ["렌더링 비용", "논리 이동과 표현 분리, 공간 배치", "VirtualConveyorBeltRenderer"],
            ["월드 규모", "청크 스트리밍 + 비활성 영역 시뮬레이션", "BlockStateStore / BackgroundSimulator"],
            ["지속성", "버전형 바이너리 직렬화와 명시적 복원", "SaveGameBinarySerializer"],
        ],
        [1.1, 2.4, 2.5],
    )

    # Page 3: Architecture
    page_break(doc)
    add_heading(doc, "전체 아키텍처", 1, "02 / ARCHITECTURE")
    p = doc.add_paragraph(style="Lead")
    p.add_run("ProjectF의 중심은 ‘라이브 오브젝트’가 아니라 ‘권위 상태’다. 화면과 가까운 영역만 라이브로 유지하고, 나머지는 저장 상태를 기반으로 진행시킨다.")
    add_picture(doc, ARCH_IMG, "그림 1. 입력, 라이브 런타임, 권위 상태, 백그라운드 시뮬레이션, 표현 계층의 관계", width=5.7)
    add_heading(doc, "아키텍처 원칙", 2)
    add_bullet(doc, "단일 권위 상태: 청크 좌표와 설치물 앵커를 기준으로 논리 상태를 저장해 중복 소유를 줄인다.", bullet_num, "단일 권위 상태:")
    add_bullet(doc, "Live / Virtual 분리: 로드된 영역은 MonoBehaviour가 담당하고, 비활성 영역은 직렬화 가능한 상태로 계산한다.", bullet_num, "Live / Virtual 분리:")
    add_bullet(doc, "스케줄 기반 실행: 모든 객체가 임의의 Update를 갖는 대신, 간격별 버킷과 쿼터로 작업을 분산한다.", bullet_num, "스케줄 기반 실행:")
    add_bullet(doc, "표현의 후행성: 렌더링은 논리 상태의 결과를 소비하며, GPU 경로 실패 시에도 기능이 유지되는 대체 경로를 둔다.", bullet_num, "표현의 후행성:")
    add_callout(doc, "확장성의 기준", "객체 수가 증가할 때 비용이 ‘컴포넌트 수 × 프레임’으로 선형 누적되지 않도록, 상태·업데이트·표현의 수명주기를 각각 통제한다.")

    # Page 4: Conveyor
    page_break(doc)
    add_heading(doc, "컨베이어 시뮬레이션과 렌더링 최적화", 1, "03 / CONVEYOR")
    p = doc.add_paragraph(style="Lead")
    p.add_run("팩토리오의 컨베이어 최적화 방식을 연구하고 AI를 설계 보조 도구로 활용해, 대량 아이템 이동을 개별 GameObject 중심 구조에서 데이터·배치 중심 구조로 옮겼다.")
    add_picture(doc, CONVEYOR_IMG, "그림 2. 컨베이어의 논리 업데이트부터 GPU 표현까지 이어지는 처리 단계", width=5.3)
    add_heading(doc, "구현 포인트", 2)
    add_bullet(doc, "MapObjectTickManager는 간격별 버킷과 쿼터·커서로 Tick 부하를 여러 프레임에 분산한다.", bullet_num)
    add_bullet(doc, "TransformJobProcessor는 Persistent NativeArray를 재사용하고 IJobParallelFor로 행렬·셀 좌표를 계산한다.", bullet_num)
    add_bullet(doc, "VirtualConveyorBeltRenderer는 렌더 키와 공간 셀 기준으로 인스턴스를 배치한다.", bullet_num)
    add_bullet(doc, "BRG 백엔드는 GPU 버퍼를 사용하고, 셰이더·플랫폼이 맞지 않으면 기존 렌더러로 복귀한다.", bullet_num)
    add_callout(
        doc,
        "구조적 결과",
        "아이템 이동과 시각 오브젝트의 생명주기를 분리해 CPU Transform 계산과 GPU 제출을 배치 단위로 다룬다. Diagnostics는 라이브/저장 아이템 수, 중복 상주, 잘못된 레인·연결을 점검한다. 동일 조건의 전·후 벤치마크가 없어 FPS·Draw Call 개선률은 기재하지 않았다.",
        accent=SUCCESS,
    )

    # Page 5: Railway
    page_break(doc)
    add_heading(doc, "철도 경로 탐색과 자동 운행", 1, "04 / RAILWAY")
    p = doc.add_paragraph(style="Lead")
    p.add_run("철도는 단순한 선형 이동이 아니라, 연결 그래프·열차 편성·역 정차·연료·선로 점유가 동시에 영향을 주는 상태 기반 시스템으로 설계했다.")
    add_picture(doc, RAIL_IMG, "그림 3. 선로 샘플링, 그래프 구성, 경로 탐색, 자동 운행 상태 머신", width=5.75)
    add_heading(doc, "경로 데이터", 2)
    add_bullet(doc, "Railload는 렌더된 중심선 샘플과 누적 거리를 캐시해 위치와 접선을 반복 계산하지 않는다.", bullet_num)
    add_bullet(doc, "AutoDriveRoutePlanner는 선로 연결을 노드와 간선으로 구성하고, 우선순위 큐 기반 탐색에 역주행·급회전 비용을 반영한다.", bullet_num)
    add_bullet(doc, "연결 그래프 버전이 변하면 관련 캐시를 무효화해 과거 노선의 경로를 재사용하지 않는다.", bullet_num)
    add_heading(doc, "운행 상태와 복원", 2)
    add_bullet(doc, "Idle, Planning, Moving, Docking, WaitingAtStation, WaitingForFuel, WaitingForPath, WaitingForClearTrack, Arrived 상태를 명시적으로 분리한다.", bullet_num)
    add_bullet(doc, "연결된 차량은 하나의 편성 그래프로 수집하고, 대표 컨트롤러와 연결 리비전을 통해 구성을 추적한다.", bullet_num)
    add_bullet(doc, "현재 선로 샘플, 목적지, 자동 운행 상태는 저장 데이터에 포함되어 청크 복원 이후에도 운행 맥락을 이어간다.", bullet_num)
    add_callout(doc, "설계 효과", "경로 계산, 실제 이동, 정차·대기 조건을 분리해 디버깅 지점을 명확히 했고, 선로 변경 시 캐시 갱신 범위를 추적할 수 있게 했다.")

    # Page 6: Streaming
    page_break(doc)
    add_heading(doc, "청크 스트리밍과 백그라운드 시뮬레이션", 1, "05 / WORLD STREAMING")
    p = doc.add_paragraph(style="Lead")
    p.add_run("플레이어 주변의 세부 오브젝트만 유지하면서도 멀리 있는 공장이 멈추지 않도록, 청크 수명주기와 시뮬레이션 수명주기를 분리했다.")
    add_heading(doc, "언로드부터 복원까지", 2)
    add_numbered(doc, "TerrainChunkStreamingScheduler가 생성·언로드 요청을 큐와 HashSet으로 중복 제거하고, 프레임별 예산으로 처리한다.", streaming_num)
    add_numbered(doc, "언로드 시 블록·자원·설치물·컨베이어·열차 상태를 BlockStateStore에 캡처한다.", streaming_num)
    add_numbered(doc, "비활성 청크는 InstallationBackgroundSimulator가 직렬화된 상태를 대상으로 생산·입출력·운송 시간을 진행한다.", streaming_num)
    add_numbered(doc, "로드 시 의존성이 큰 설치물부터 순서와 예산을 적용해 복원하고, 컨베이어 아이템·열차 샘플·자동 운행 상태를 다시 연결한다.", streaming_num)
    add_heading(doc, "안정성을 위한 경계", 2)
    add_matrix_table(
        doc,
        ["경계", "보호 장치", "의도"],
        [
            ["로드/언로드 흔들림", "서로 다른 로드·언로드 반경", "경계에서 반복 생성 방지"],
            ["중복 요청", "Queue + HashSet", "같은 청크의 중복 처리 방지"],
            ["프레임 스파이크", "코루틴과 프레임별 복원 예산", "작업량 분산"],
            ["상태 중복", "라이브/저장 상주 진단", "아이템 이중 소유 탐지"],
        ],
        [1.3, 2.1, 2.6],
    )
    add_callout(doc, "핵심 불변식", "한 시점의 아이템·설치물 상태는 라이브 월드와 저장 상태 중 하나가 권위를 가져야 한다. 스트리밍은 이 권위를 넘기는 과정이며, 진단 코드는 중복 상주와 연결 손상을 검사한다.", accent=SUCCESS)

    # Page 7: Save and data
    page_break(doc)
    add_heading(doc, "저장 포맷과 데이터 원본", 1, "06 / PERSISTENCE & DATA")
    p = doc.add_paragraph(style="Lead")
    p.add_run("저장은 월드의 논리 상태를 다시 구성하기 위한 계약으로 다룬다.")
    add_heading(doc, "버전형 바이너리 저장", 2)
    add_bullet(doc, "PF_SAVE 매직 헤더와 버전 번호를 먼저 기록하고, 지원 범위를 벗어난 파일을 거부한다.", bullet_num)
    add_bullet(doc, "임시 파일 쓰기를 완료한 뒤 기존 파일을 교체해 미완성 데이터의 노출 시간을 줄인다.", bullet_num)
    add_bullet(doc, "지형·설치물·컨베이어·열차·유체·플레이어 상태를 명시적으로 직렬화한다.", bullet_num)
    add_heading(doc, "데이터 진실 원본", 2)
    add_matrix_table(
        doc,
        ["데이터", "진실 원본", "런타임 소비 방식"],
        [
            ["아이템", "ItemDefinition ScriptableObject", "ID·에너지·유체·제작 속성 참조"],
            ["제작 트리", "crafting_tree.bytes 바이너리", "CraftingTreeRuntime이 역색인 Dictionary 구성"],
            ["월드 상태", "BlockStateStore + SaveGameData", "로드 범위에 따라 라이브 오브젝트로 복원"],
        ],
        [1.2, 2.0, 2.8],
    )
    add_heading(doc, "유체 네트워크", 2)
    p = doc.add_paragraph()
    p.add_run("파이프와 탱크는 연결 마스크·네트워크 탐색을 사용하고, 토폴로지 버전과 캐시로 재탐색을 줄인다. 지하 파이프는 원격 쌍을, 펌프는 출력망의 용량·충전 비율을 관리한다.")
    add_callout(doc, "데이터 설계의 목적", "편집 원본, 런타임 산출물, 저장 상태를 구분해 동기화 방향과 포맷 변경의 복원 영향 범위를 드러낸다.")

    # Page 8: Diagnostics & AI
    page_break(doc)
    add_heading(doc, "진단 도구와 AI 활용 개발", 1, "07 / TOOLING & AI")
    p = doc.add_paragraph(style="Lead")
    p.add_run("AI는 코드를 대신 책임지는 도구가 아니라, 복잡한 최적화 대안을 빠르게 탐색하고 구현 범위를 넓히는 협업 도구로 활용했다.")
    add_heading(doc, "검증을 위한 하네스", 2)
    add_bullet(doc, "ConveyorRuntimeDiagnostics는 로드·저장 블록의 아이템 총량, 중복 상주, 잘못된 레인·링크를 감사한다.", bullet_num)
    add_bullet(doc, "MapObjectTickManager의 ProfilerMarker·Counter는 Tick 비용과 분산 상태를 관찰할 수 있는 진입점을 제공한다.", bullet_num)
    add_bullet(doc, "별도 WinForms 기반 MapObjectProfilerTool은 localhost TCP로 FPS, 프레임 시간, Tick 루프, 컨베이어 상태 지표를 수집한다.", bullet_num)
    add_heading(doc, "AI 활용 루프", 2)
    add_numbered(doc, "문제 정의: ‘오브젝트 수가 늘수록 어떤 비용이 누적되는가’를 상태·연산·표현으로 분해한다.", ai_num)
    add_numbered(doc, "대안 탐색: 팩토리오식 데이터 중심 처리, Unity Jobs/Burst, 공간 배치, BRG 적용 가능성을 AI와 비교한다.", ai_num)
    add_numbered(doc, "구현·리뷰: AI가 제안한 구조를 프로젝트의 권위 상태, 청크 수명주기, GC 제약에 맞게 수정하고 중복 코드를 제거한다.", ai_num)
    add_numbered(doc, "검증: 진단 API와 프로파일러로 상태 불변식과 비용을 확인하고, 실패 가능한 GPU 경로에는 폴백을 둔다.", ai_num)
    add_callout(
        doc,
        "AI 활용으로 확장한 구현 범위",
        "혼자서 진입하기 어려웠던 Burst Job, BatchRendererGroup, 그래프 기반 경로 탐색, 백그라운드 상태 시뮬레이션을 작은 검증 단위로 나누어 구현할 수 있었다. 최종 설계 선택과 코드 통합, 오류 판단은 개발자가 직접 수행했다.",
        accent=BLUE,
    )

    # Page 9: code map / next
    page_break(doc)
    add_heading(doc, "코드 맵과 다음 검증 과제", 1, "08 / CODE MAP")
    p = doc.add_paragraph(style="Lead")
    p.add_run("핵심 시스템은 역할별 파일로 분리되어 있으며, 아래 경로가 기술 검토의 시작점이다.")
    add_matrix_table(
        doc,
        ["영역", "대표 코드", "검토 포인트"],
        [
            ["컨베이어·렌더", "VirtualConveyorBeltRenderer.cs\nTransformJobProcessor.cs / BRGBackend.cs", "배치 키·버퍼·폴백"],
            ["Tick", "MapObjectTickManager.cs", "버킷·쿼터·Profiler"],
            ["철도", "SteamTrain.cs / Train.cs / Railload.cs", "그래프·경로·상태 머신"],
            ["월드·스트리밍", "BlockStateStore.cs / BackgroundSimulator.cs\nChunkStreamingScheduler.cs", "권위 상태·복원 예산"],
            ["저장", "SaveGameBinarySerializer.cs\nSaveGameData.cs", "포맷 버전·복원 계약"],
            ["진단", "ConveyorRuntimeDiagnostics.cs\nTools/MapObjectProfilerTool", "상태 감사·성능 지표"],
        ],
        [1.05, 2.65, 2.3],
    )
    add_heading(doc, "다음 우선순위", 2)
    add_numbered(doc, "동일 맵·아이템 수를 고정한 자동 벤치마크를 만들어 CPU 시간, 배치 수, GC Alloc을 전·후 비교한다.", next_num)
    add_numbered(doc, "SteamTrain의 경로 계획·운행 상태·편성 제어 책임을 점진적으로 분리해 테스트 가능한 경계를 넓힌다.", next_num)
    add_numbered(doc, "컨베이어와 스트리밍 시나리오를 재현하는 테스트 하네스를 정식화하고 CI에서 저장·복원 불변식을 검사한다.", next_num)
    add_numbered(doc, "의존 관계가 안정된 시스템부터 namespace와 asmdef를 도입해 컴파일 경계를 명확히 한다.", next_num)
    add_callout(doc, "정리", "ProjectF의 기술적 핵심은 기능의 수가 아니라, 대규모 월드에서 상태의 권위와 비용의 발생 지점을 통제하는 구조다. 컨베이어·철도·스트리밍·저장은 이 원칙을 서로 다른 문제에 적용한 결과다.", accent=SUCCESS)

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    result = build_document()
    print(result)
