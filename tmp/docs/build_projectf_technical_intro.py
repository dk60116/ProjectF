from pathlib import Path
from datetime import date

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Git\ProjectF")
TMP = ROOT / "tmp" / "docs"
DIAGRAMS = TMP / "diagrams"
OUTPUT = ROOT / "output" / "docx" / "ProjectF_기술_소개서.docx"
FONT_PATH = Path(r"C:\Windows\Fonts\malgun.ttf")
FONT_BOLD_PATH = Path(r"C:\Windows\Fonts\malgunbd.ttf")

NAVY = "17324D"
BLUE = "2E74B5"
LIGHT_BLUE = "E8F1F8"
PALE = "F4F7FA"
GOLD = "D9A441"
CHARCOAL = "263746"
MUTED = "607485"
WHITE = "FFFFFF"
LINE = "CBD8E3"
GREEN = "3C8D78"
RED = "C45D52"


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def pil(hex_value):
    return hex_value if hex_value.startswith("#") else "#" + hex_value


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + edge
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    total = sum(widths)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        existing = tbl_pr.find(qn(tag))
        if existing is not None:
            tbl_pr.remove(existing)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, size=None, bold=None, color=None, name="맑은 고딕"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = rgb(color)


def set_paragraph_border(paragraph, color=BLUE, size=16, space=6, side="left"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    edge = OxmlElement("w:" + side)
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), str(space))
    edge.set(qn("w:color"), color)
    p_bdr.append(edge)


def set_paragraph_shading(paragraph, fill=PALE):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("PROJECTF  ·  ")
    set_run_font(run, 8, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run2 = paragraph.add_run()
    set_run_font(run2, 8, color=MUTED)
    run2._r.extend([fld_char1, instr_text, fld_char2])


def create_bullet_num(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.extend([tabs, ind])
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    r_pr.append(fonts)
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr, r_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_bullet(doc, text, num_id, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.12
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    p_pr.append(num_pr)
    if bold_prefix and text.startswith(bold_prefix):
        lead = p.add_run(bold_prefix)
        set_run_font(lead, 10.3, bold=True, color=CHARCOAL)
        body = p.add_run(text[len(bold_prefix):])
        set_run_font(body, 10.3, color=CHARCOAL)
    else:
        run = p.add_run(text)
        set_run_font(run, 10.3, color=CHARCOAL)
    return p


def add_body(doc, text, size=10.3, color=CHARCOAL, after=6, line=1.22, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    run = p.add_run(text)
    set_run_font(run, size, bold=bold, color=color)
    return p


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.upper())
    set_run_font(run, 8.5, bold=True, color=GOLD)
    run.font.letter_spacing = Pt(0.8) if hasattr(run.font, "letter_spacing") else None
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    return p


def add_callout(doc, title, text, color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.05)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    set_paragraph_border(p, color=color)
    set_paragraph_shading(p, fill=PALE)
    r1 = p.add_run(title + "  ")
    set_run_font(r1, 10, bold=True, color=color)
    r2 = p.add_run(text)
    set_run_font(r2, 9.8, color=CHARCOAL)
    return p


def add_source(doc, paths):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run("CODE ANCHORS  " + "  ·  ".join(paths))
    set_run_font(r, 7.4, color=MUTED, name="Consolas")
    return p


def add_figure(doc, path, width=6.45, alt_text=""):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inline_shape = p.add_run().add_picture(str(path), width=Inches(width))
    if alt_text:
        doc_pr = inline_shape._inline.docPr
        doc_pr.set("descr", alt_text)
        doc_pr.set("title", alt_text)
    return p


def insert_page_break(doc):
    paragraph = doc.paragraphs[-1]
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(header)
        set_run_font(r, 9.1, bold=True, color=WHITE)
    header_tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "1")
    header_tr_pr.append(tbl_header)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            if row_index % 2 == 0:
                set_cell_shading(cells[idx], "F7F9FB")
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(value)
            set_run_font(r, 8.65, color=CHARCOAL)
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1)
    return table


def font(size, bold=False):
    path = FONT_BOLD_PATH if bold and FONT_BOLD_PATH.exists() else FONT_PATH
    return ImageFont.truetype(str(path), size)


def arrow(draw, a, b, color=BLUE, width=5):
    draw.line([a, b], fill=pil(color), width=width)
    x1, y1 = a
    x2, y2 = b
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        tip = [(x2, y2), (x2 - direction * 18, y2 - 11), (x2 - direction * 18, y2 + 11)]
    else:
        direction = 1 if y2 > y1 else -1
        tip = [(x2, y2), (x2 - 11, y2 - direction * 18), (x2 + 11, y2 - direction * 18)]
    draw.polygon(tip, fill=pil(color))


def centered(draw, xy, text, fnt, fill=CHARCOAL, max_width=None):
    x, y = xy
    lines = text.split("\n")
    line_h = fnt.size + 8
    total_h = line_h * len(lines)
    for i, line in enumerate(lines):
        box = draw.textbbox((0, 0), line, font=fnt)
        w = box[2] - box[0]
        draw.text((x - w / 2, y - total_h / 2 + i * line_h), line, font=fnt, fill=pil(fill))


def box(draw, rect, title, subtitle="", fill=WHITE, outline=LINE, accent=BLUE):
    draw.rounded_rectangle(rect, radius=18, fill=pil(fill), outline=pil(outline), width=3)
    x1, y1, x2, y2 = rect
    draw.rounded_rectangle((x1, y1, x1 + 10, y2), radius=5, fill=pil(accent))
    center_y = (y1 + y2) / 2
    centered(draw, ((x1 + x2) / 2 + 5, center_y - (15 if subtitle else 0)), title, font(29, True), NAVY)
    if subtitle:
        centered(draw, ((x1 + x2) / 2 + 5, center_y + 28), subtitle, font(21), MUTED)


def save_architecture_diagram(path):
    img = Image.new("RGB", (1600, 520), pil(PALE))
    draw = ImageDraw.Draw(img)
    layers = [
        ("입력 · UI", "건설 / 조작 / 상태 표시", GOLD),
        ("설치물 시뮬레이션", "컨베이어 · 생산 · 철도 · 유체", BLUE),
        ("월드 · 청크", "생성 / 로드 / 언로드 스케줄링", GREEN),
        ("상태 · 영속화", "BlockStateStore / Binary Save", RED),
        ("렌더링 · 진단", "BRG / Instancing / Profiler", NAVY),
    ]
    y = 46
    for i, (title, subtitle, accent) in enumerate(layers):
        rect = (110, y, 1490, y + 72)
        box(draw, rect, title, subtitle, fill=WHITE, outline=LINE, accent=accent)
        if i < len(layers) - 1:
            arrow(draw, (800, y + 76), (800, y + 91), color=MUTED, width=4)
        y += 91
    img.save(path)


def save_conveyor_diagram(path):
    img = Image.new("RGB", (1700, 500), pil(PALE))
    draw = ImageDraw.Draw(img)
    stages = [
        ("논리 상태", "lane / link", GOLD),
        ("거주지 판정", "live ↔ saved", RED),
        ("변환 계산", "Burst Job", GREEN),
        ("공간 배칭", "mesh · material · cell", BLUE),
        ("렌더 백엔드", "BRG / fallback", NAVY),
        ("진단", "count · duplicate", RED),
    ]
    gap = 32
    w = 245
    x = 38
    for idx, (title, subtitle, accent) in enumerate(stages):
        rect = (x, 112, x + w, 300)
        box(draw, rect, title, subtitle, fill=WHITE, outline=LINE, accent=accent)
        if idx < len(stages) - 1:
            arrow(draw, (x + w + 5, 206), (x + w + gap - 5, 206), color=BLUE, width=5)
        x += w + gap
    centered(draw, (850, 390), "프레임 비용과 상태 정합성을 같은 파이프라인에서 관찰", font(27, True), NAVY)
    centered(draw, (850, 432), "측정 가능한 카운터를 남겨 최적화 효과를 검증 가능한 형태로 설계", font(22), MUTED)
    img.save(path)


def save_rail_diagram(path):
    img = Image.new("RGB", (1600, 520), pil(PALE))
    draw = ImageDraw.Draw(img)
    nodes = {
        "A": (150, 250), "N1": (420, 250), "N2": (690, 135),
        "N3": (690, 365), "N4": (980, 250), "B": (1435, 250),
    }
    edges = [("A", "N1"), ("N1", "N2"), ("N1", "N3"), ("N2", "N4"), ("N3", "N4"), ("N4", "B")]
    for a, b in edges:
        draw.line([nodes[a], nodes[b]], fill=pil("AABBC8"), width=13)
    route = [("A", "N1"), ("N1", "N2"), ("N2", "N4"), ("N4", "B")]
    for a, b in route:
        draw.line([nodes[a], nodes[b]], fill=pil(BLUE), width=9)
    for name, pos in nodes.items():
        radius = 56 if name in ("A", "B") else 31
        fill = GOLD if name in ("A", "B") else WHITE
        draw.ellipse((pos[0]-radius, pos[1]-radius, pos[0]+radius, pos[1]+radius), fill=pil(fill), outline=pil(NAVY), width=5)
        centered(draw, pos, "역 " + name if name in ("A", "B") else name, font(27, True), NAVY)
    centered(draw, (800, 55), "샘플링된 레일 경로 → 연결 그래프 → 비용 기반 경로 탐색", font(31, True), NAVY)
    centered(draw, (800, 465), "기본 거리 + 후진 출발 · 후진 전환 · 급커브 패널티", font(24), MUTED)
    img.save(path)


def save_persistence_diagram(path):
    img = Image.new("RGB", (1700, 510), pil(PALE))
    draw = ImageDraw.Draw(img)
    upper = [
        ("활성 월드", "GameObject 상태", BLUE),
        ("캡처", "ID · 설치물 · 아이템", GOLD),
        ("상태 저장", "BlockStateStore / binary", RED),
        ("백그라운드", "비활성 청크 시뮬레이션", GREEN),
        ("복원", "청크 생성 후 재결합", NAVY),
    ]
    gap = 42
    w = 280
    x = 55
    for idx, (title, subtitle, accent) in enumerate(upper):
        rect = (x, 58, x + w, 220)
        box(draw, rect, title, subtitle, fill=WHITE, outline=LINE, accent=accent)
        if idx < len(upper) - 1:
            arrow(draw, (x + w + 6, 139), (x + w + gap - 6, 139), color=BLUE, width=5)
        x += w + gap
    centered(draw, (850, 302), "AI 활용 루프", font(28, True), GOLD)
    steps = ["문제 분해", "대안 탐색", "로컬 구현", "진단·프로파일링", "피드백·리팩터링"]
    x = 170
    for idx, step in enumerate(steps):
        centered(draw, (x, 390), step, font(23, True), NAVY)
        if idx < len(steps) - 1:
            arrow(draw, (x + 83, 390), (x + 180, 390), color=GOLD, width=4)
        x += 320
    img.save(path)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(10.3)
    normal.font.color.rgb = rgb(CHARCOAL)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.22

    for level, size, color, before, after in ((1, 16, BLUE, 8, 8), (2, 12.5, BLUE, 8, 5), (3, 11, NAVY, 6, 3)):
        style = styles[f"Heading {level}"]
        style.font.name = "맑은 고딕"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_document_settings(doc):
    settings = doc.settings._element
    even_odd = settings.find(qn("w:evenAndOddHeaders"))
    if even_odd is not None:
        settings.remove(even_odd)


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.88)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.34)
    section.footer_distance = Inches(0.35)
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("PROJECTF  /  TECHNICAL INTRODUCTION")
    set_run_font(r, 7.7, bold=True, color=MUTED)
    add_page_number(section.footer.paragraphs[0])


def build_document():
    DIAGRAMS.mkdir(parents=True, exist_ok=True)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    architecture = DIAGRAMS / "architecture.png"
    conveyor = DIAGRAMS / "conveyor_pipeline.png"
    rail = DIAGRAMS / "rail_graph.png"
    persistence = DIAGRAMS / "persistence_ai.png"
    save_architecture_diagram(architecture)
    save_conveyor_diagram(conveyor)
    save_rail_diagram(rail)
    save_persistence_diagram(persistence)

    doc = Document()
    configure_styles(doc)
    configure_document_settings(doc)
    configure_section(doc.sections[0])
    doc.core_properties.title = "ProjectF 기술 소개서"
    doc.core_properties.subject = "Unity 기반 자동화·생존 시뮬레이션의 시스템 설계"
    doc.core_properties.author = "한택근"
    doc.core_properties.keywords = "Unity, C#, ProjectF, Conveyor, Railway, Save Load, AI"
    bullet_id = create_bullet_num(doc)

    # PAGE 1 — COVER
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(44)
    add_kicker(doc, "TECHNICAL PORTFOLIO · 2026")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("ProjectF")
    set_run_font(r, 35, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(25)
    r = p.add_run("Unity 기반 자동화·생존 시뮬레이션의 시스템 설계")
    set_run_font(r, 17, bold=True, color=BLUE)

    add_callout(
        doc,
        "PROJECT DEFINITION",
        "자원 채집부터 생산 설비, 컨베이어 물류, 철도 운송, 유체 네트워크까지 하나의 월드에서 연결되는 1인 개발 자동화 게임입니다. 확장되는 공장의 처리량과 월드 상태의 정확성을 동시에 유지하는 것을 핵심 기술 목표로 삼았습니다.",
        color=GOLD,
    )

    add_heading(doc, "기술 스냅샷", 2)
    add_bullet(doc, "엔진: Unity 6000.4.0f1 · 렌더 파이프라인: URP 17.4", bullet_id)
    add_bullet(doc, "언어·런타임: C# · Burst 1.8.28 · Collections 6.4.0 · Mathematics 1.3.3", bullet_id)
    add_bullet(doc, "핵심 시스템: 컨베이어, 생산 설치물, 철도·차량, 파이프·유체, 청크 스트리밍, 세이브/로드", bullet_id)
    add_bullet(doc, "개발 방식: 성능 민감 경로를 측정 가능한 구조로 만들고 AI를 분석·구현 보조 도구로 활용", bullet_id)

    add_heading(doc, "설계 원칙", 2)
    add_body(doc, "① 논리 상태와 표현 계층을 분리합니다.  ② 활성·비활성 청크의 상태를 하나의 권위 모델로 연결합니다.  ③ 고비용 처리는 배치화하고 실패 시 안전한 대체 경로를 둡니다.  ④ 프로파일러와 진단 지표로 최적화 결과를 검증할 수 있게 만듭니다.", after=10)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("한택근  |  Client Programmer")
    set_run_font(r, 11, bold=True, color=NAVY)
    r2 = p.add_run("\nProjectF v0.1  ·  Repository-based technical brief")
    set_run_font(r2, 8.7, color=MUTED)

    insert_page_break(doc)

    # PAGE 2 — ARCHITECTURE
    add_kicker(doc, "01 · SYSTEM ARCHITECTURE")
    add_heading(doc, "확장 가능한 공장을 위한 계층형 구조", 1)
    add_body(doc, "ProjectF의 핵심은 ‘월드에 존재하는 모든 객체를 항상 실행하지 않는다’는 판단입니다. 입력과 설치물 로직, 청크 수명주기, 영속 상태, 렌더링·진단을 분리하여 각 계층의 비용과 책임을 통제합니다.", after=5)
    add_figure(doc, architecture, 6.42, "입력과 UI부터 설치물 시뮬레이션, 월드 청크, 상태 영속화, 렌더링과 진단으로 이어지는 ProjectF 계층 구조")

    add_heading(doc, "청크 스트리밍과 상태 수명주기", 2)
    add_bullet(doc, "TerrainChunkStreamingScheduler는 생성·언로드 요청을 큐와 HashSet으로 중복 제거하고, 코루틴 예산에 맞춰 순차 처리합니다.", bullet_id)
    add_bullet(doc, "청크 언로드 시 설치물과 운반 아이템을 BlockStateStore로 캡처하고, 재진입 시 지형 생성 후 제한된 개수만 단계적으로 복원합니다.", bullet_id)
    add_bullet(doc, "화면 밖 청크는 InstallationBackgroundSimulator가 저장 상태를 기준으로 진행시켜, 비활성 영역도 생산 흐름이 단절되지 않게 합니다.", bullet_id)

    add_heading(doc, "논리 상태와 표현의 분리", 2)
    add_body(doc, "GameObject의 존재 여부를 게임 상태의 진실 원본으로 삼지 않습니다. 활성 청크의 런타임 객체와 비활성 청크의 저장 상태가 동일한 식별자·연결 정보를 공유하고, 렌더러는 그 결과를 소비하는 표현 계층으로 동작합니다. 이 구조는 대규모 월드에서 객체 수를 줄이면서도 저장·복원 정합성을 유지하기 위한 기반입니다.", after=5)
    add_callout(doc, "핵심 효과", "청크 경계를 넘나드는 생산·운송 상태를 보존하면서, 프레임당 생성·복원 작업량을 제한할 수 있습니다.")
    add_source(doc, ["TerrainChunkStreamingScheduler.cs", "TerrainGenerator.ChunkPersistence.cs", "BlockStateStore.cs", "InstallationBackgroundSimulator.cs"])

    insert_page_break(doc)

    # PAGE 3 — CONVEYOR
    add_kicker(doc, "02 · CONVEYOR OPTIMIZATION")
    add_heading(doc, "컨베이어를 ‘개별 오브젝트’가 아닌 데이터 흐름으로 처리", 1)
    add_body(doc, "공장 규모가 커질수록 벨트와 운반 아이템은 CPU 갱신, Transform 계산, 드로우콜, GC 할당을 동시에 압박합니다. 이를 해결하기 위해 논리 상태 → 변환 계산 → 공간 배칭 → 렌더링 → 진단으로 이어지는 파이프라인을 구성했습니다.", after=4)
    add_figure(doc, conveyor, 6.45, "컨베이어 논리 상태에서 거주지 판정, Burst 변환 계산, 공간 배칭, 렌더 백엔드, 진단으로 이어지는 처리 파이프라인")

    add_table(
        doc,
        ["병목", "구현 방식", "검증 지점"],
        [
            ["CPU·GC", "Persistent NativeArray와 Burst IJobParallelFor로 아이템 변환을 일괄 계산", "작업 수·프레임 시간·할당"],
            ["드로우콜", "메시·머티리얼·공간 셀 기준 배치, BatchRendererGroup 사용", "등록 벨트·예상 드로우콜"],
            ["호환성", "BRG 지원 여부를 확인하고 RenderMeshInstanced로 안전하게 폴백", "백엔드 실패·전환 로그"],
            ["상태 오류", "활성·저장 아이템을 함께 집계하고 중복 거주·잘못된 링크 탐지", "live/saved/duplicate 카운터"],
        ],
        [1700, 4700, 2960],
    )

    add_heading(doc, "최적화의 기준", 2)
    add_body(doc, "작은 작업은 메인 스레드에서 처리해 잡 스케줄링 오버헤드를 피하고, 임계값을 넘으면 Burst 병렬 경로로 전환합니다. 버퍼 용량은 2의 거듭제곱으로 확장해 재할당 빈도를 낮추고, 렌더 백엔드는 기능 지원과 셰이더 호환성을 확인한 뒤 선택합니다.", after=4)
    add_callout(doc, "검증 원칙", "현재 저장소에는 정량 벤치마크 결과 대신 프로파일러·진단 카운터가 준비되어 있습니다. 따라서 성능 향상 수치를 과장하지 않고, 같은 맵·같은 아이템 수에서 반복 측정할 수 있는 구조적 기반을 성과로 제시합니다.", color=GREEN)
    insert_page_break(doc)

    # PAGE 4 — RAIL / FLUID
    add_kicker(doc, "03 · NETWORK SIMULATION")
    add_heading(doc, "철도 자동운전: 곡선을 그래프로 바꾸고 비용으로 판단", 1)
    add_body(doc, "레일은 단순한 직선 연결이 아니라 샘플링된 곡선과 분기점의 집합입니다. 각 레일의 렌더 경로를 노드·엣지 그래프로 변환하고, 역을 출발지와 목적지로 연결해 비용 기반 경로를 탐색합니다.", after=4)
    add_figure(doc, rail, 6.38, "역 A와 역 B 사이의 레일 곡선을 연결 그래프로 변환하고 비용이 낮은 경로를 선택하는 예시")

    add_heading(doc, "경로 계획과 주행 상태", 2)
    add_bullet(doc, "AutoDriveRoutePlanner는 그래프 버전을 캐시하고, 레일 연결 변경 시에만 무효화하여 불필요한 재구성을 줄입니다.", bullet_id)
    add_bullet(doc, "거리뿐 아니라 후진 출발, 진행 방향 전환, 급격한 회전에 패널티를 부여해 실제 주행에 가까운 경로를 선택합니다.", bullet_id)
    add_bullet(doc, "Idle → Planning → Moving → Docking → WaitingAtStation/Arrived 상태를 분리하고, 연료·경로·선로 점유 대기 원인을 별도 상태로 노출합니다.", bullet_id)
    add_bullet(doc, "열차 연결 그래프를 순회해 편성 전체를 구성하고, 경로 세그먼트·커서·look-ahead를 저장해 주행을 이어갑니다.", bullet_id)

    add_heading(doc, "같은 원칙을 유체 네트워크에 적용", 2)
    add_body(doc, "Pipe·UndergroundPipe·FluidTank·Pump는 연결 토폴로지를 BFS로 탐색하고 버전 단위로 캐시합니다. 지하 파이프의 원격 링크까지 같은 네트워크로 묶어 유체량을 평준화하되, 연결이 바뀔 때만 네트워크를 다시 계산합니다. 철도와 유체는 모두 ‘연결 관계를 그래프로 명시하고 변경 시점에만 갱신한다’는 공통 설계를 사용합니다.", after=4)
    add_source(doc, ["SteamTrain.cs", "Train.cs", "Railload.cs", "RailConnectionUtility.cs", "Pipe.cs", "UndergroundPipe.cs"])

    insert_page_break(doc)

    # PAGE 5 — PERSISTENCE / VALIDATION / AI
    add_kicker(doc, "04 · PERSISTENCE, VALIDATION & AI")
    add_heading(doc, "세이브·복원과 AI 활용을 하나의 검증 루프로", 1)
    add_body(doc, "세이브 시스템은 설치물 타입별 상태, 위치·회전, 내부 아이템, 연결 정보, 열차 자동운전 상태를 바이너리 포맷으로 직렬화합니다. 청크 저장 상태와 전체 세이브가 동일한 복원 규칙을 공유하도록 하여, 스트리밍과 재접속에서 다른 결과가 생기지 않게 설계했습니다.", after=4)
    add_figure(doc, persistence, 6.45, "활성 월드의 상태 캡처와 저장, 백그라운드 시뮬레이션, 복원 과정 및 AI 활용 반복 루프")

    add_heading(doc, "AI를 복잡한 구현의 가속기로 사용", 2)
    add_body(doc, "Factorio의 컨베이어 최적화 방식을 공부한 뒤, AI와 함께 오브젝트 단위 갱신의 병목을 데이터 중심 파이프라인으로 재구성했습니다. AI는 대형 코드베이스 탐색, BRG·Burst처럼 익숙하지 않은 API의 대안 비교, 반복되는 로직의 리팩터링 초안에 활용했습니다. 저는 상태의 진실 원본, 폴백 조건, 성능 임계값을 정의하고 생성된 코드를 직접 검토했으며, 진단 결과를 다시 입력해 중복·예외 분기를 정리했습니다. 이를 통해 혼자서는 진입 비용이 컸던 렌더 배칭과 병렬 변환 로직을 실제 프로젝트 구조 안에 구현할 수 있었습니다.", after=5)

    add_heading(doc, "검증 도구와 다음 단계", 2)
    add_bullet(doc, "MapObjectProfilerTool: TCP로 FPS·프레임 시간·틱 카운터를 조회하는 독립 프로파일러", bullet_id)
    add_bullet(doc, "ConveyorRuntimeDiagnostics: 활성/저장 아이템 수, 중복 거주, 레인·링크 오류를 한 화면에서 확인", bullet_id)
    add_bullet(doc, "BeltAlignmentProbe: 에디터에서 벨트 정렬과 샘플 위치를 재현하는 진단 진입점", bullet_id)
    add_bullet(doc, "다음 단계: 고정 맵 기반 회귀 벤치마크, 저장 포맷 테스트, namespace·asmdef 단위의 점진적 책임 분리", bullet_id)

    add_callout(doc, "기술적 결론", "ProjectF는 단순 기능 나열이 아니라, 대규모 자동화 월드에서 상태 정합성·성능·디버깅 가능성을 함께 다루는 시스템 프로젝트입니다.", color=GOLD)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
